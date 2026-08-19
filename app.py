"""
GraphGen Pro
============
Upload a survey spreadsheet -> auto-detect independent vs dependent variables ->
generate SPSS-style grouped bar charts for every IV x DV pair -> run a one-way
ANOVA -> export a Word document with figures, legends, an ANOVA table and an
auto-written Results & Discussion section.

100% free to run. Text is written by Groq (https://console.groq.com), which has
a free API tier — set a GROQ_API_KEY environment variable to turn it on. If no
key is set, or the Groq call fails/times out for any reason, the app silently
falls back to a built-in rule-based sentence generator so it always keeps
working with zero configuration and zero cost.
"""

import io
import os
import gc
import json
import base64
import uuid
import traceback

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from flask import Flask, request, render_template, jsonify, send_file
from scipy import stats

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Groq (free-tier AI) -----------------------------------------------------
# pip install groq. Get a free key at https://console.groq.com/keys
# Set it as an environment variable: GROQ_API_KEY=gsk_...
GROQ_MODEL = os.environ.get("GROQ_MODEL", "llama-3.3-70b-versatile")
try:
    from groq import Groq
    _groq_client = Groq(api_key=os.environ["GROQ_API_KEY"]) if os.environ.get("GROQ_API_KEY") else None
except Exception:
    _groq_client = None

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 8 * 1024 * 1024  # keep Render free-tier requests bounded

# In-memory store for the last few analyses so /download can rebuild the docx.
# Fine for a single small Render instance; swap for redis/db if you need more.
REPORT_CACHE = {}
MAX_CACHE = 2

# ---------------------------------------------------------------------------
# Column classification
# ---------------------------------------------------------------------------

IV_KEYWORDS = [
    "age", "gender", "sex", "education", "qualification", "occupation",
    "job", "income", "marital", "experience", "platform", "frequency of",
    "region", "location", "designation", "year of study", "grade",
]

LIKERT_SCALES = [
    (["strongly disagree", "disagree", "neutral", "agree", "strongly agree"], 1),
    (["strongly disagree", "disagree", "neutral", "agree", "strongly agree", "strongly"], 1),
    (["very low", "low", "moderate", "high", "very high"], 1),
    (["do not trust at all", "trust a little", "neutral", "trust"], 1),
    (["never", "occasionally", "monthly", "weekly", "daily"], 1),
    (["yes", "no", "maybe"], 1),
    (["yes", "no"], 1),
]


def _norm(v):
    return str(v).strip().lower()


def classify_columns(df):
    """Return (iv_cols, dv_cols) using simple, transparent heuristics.
    Columns with too many unique values (free-text fields, IDs, emails,
    "Other, please specify" boxes, etc.) are dropped entirely -- they can't
    be turned into a meaningful bar chart, and attempting to chart one with
    hundreds/thousands of categories is what actually blows up memory/time
    on a real spreadsheet (confirmed: a single such column can take 10+
    seconds and ~300MB to plot on its own)."""
    MAX_IV_CATEGORIES = 15
    MAX_DV_CATEGORIES = 10
    n_rows = max(len(df), 1)

    iv_cols, dv_cols = [], []
    for col in df.columns:
        lc = col.lower().strip()
        if lc == "timestamp" or lc.startswith("unnamed"):
            continue

        nunique = df[col].nunique(dropna=True)
        # Free-text/ID-like columns: skip regardless of keyword match --
        # e.g. near-one-unique-value-per-row (emails, comments, "specify
        # other") isn't a chartable categorical variable.
        if nunique <= 1 or nunique > 0.5 * n_rows:
            continue

        if any(k in lc for k in IV_KEYWORDS):
            if nunique <= MAX_IV_CATEGORIES:
                iv_cols.append(col)
            # else: looks like an IV by name but has too many distinct
            # values (e.g. a free-text "age" comment field) -- skip it.
        else:
            if nunique <= MAX_DV_CATEGORIES:
                dv_cols.append(col)
    return iv_cols, dv_cols


def build_ordinal_map(series):
    """Try to map a categorical Likert-like column to an ordered 1..N scale.
    Falls back to order-of-first-appearance if no known scale matches."""
    values = [v for v in series.dropna().unique().tolist()]
    norm_values = {_norm(v) for v in values}

    for scale, _ in LIKERT_SCALES:
        scale_norm = [s for s in scale]
        if norm_values.issubset(set(scale_norm)) and len(norm_values) >= 2:
            ordered = [s for s in scale_norm if s in norm_values]
            mapping = {}
            for v in values:
                nv = _norm(v)
                if nv in ordered:
                    mapping[v] = ordered.index(nv) + 1
            if len(mapping) == len(values):
                return mapping, True

    # fallback: alphabetical-ish, first-seen order (not a real scale)
    mapping = {v: i + 1 for i, v in enumerate(values)}
    return mapping, False


def is_numeric_series(series):
    return pd.api.types.is_numeric_dtype(series)


def to_numeric_dv(df, col):
    """Return a numeric Series for a DV column, and whether it is a real scale."""
    s = df[col]
    if is_numeric_series(s):
        return s.astype(float), True
    mapping, is_real_scale = build_ordinal_map(s.dropna())
    return s.map(mapping).astype(float), is_real_scale


# ---------------------------------------------------------------------------
# Chart generation (SPSS-style grouped bar chart, % within each IV group)
# ---------------------------------------------------------------------------

# Colors sampled directly from real SPSS default categorical output
# (blue / green / tan / ...), so generated charts read as genuine SPSS exports.
SPSS_PALETTE = [
    "#3E58AC", "#2EB848", "#D3CE97", "#8C564B",
    "#9467BD", "#17BECF", "#E377C2", "#BCBD22",
]
SPSS_BG = "#F0F0F0"


def _ordered_categories(series, prefer_likert=True):
    """Return the category order to plot in: the recognized Likert-scale order
    if one matches, otherwise first-appearance order (never alphabetical --
    that's what makes SPSS output feel 'random' when re-implemented naively)."""
    values = [v for v in series.dropna().unique().tolist()]
    if prefer_likert:
        norm_values = {_norm(v) for v in values}
        for scale, _ in LIKERT_SCALES:
            if norm_values.issubset(set(scale)) and len(norm_values) >= 2:
                ordered = [s for s in scale if s in norm_values]
                # map back to the original (non-normalized) string for each slot
                lookup = {_norm(v): v for v in values}
                return [lookup[o] for o in ordered if o in lookup]
    return values  # first-appearance order


def _wrap_label(text, width=18):
    import textwrap
    return "\n".join(textwrap.wrap(str(text), width=width)) or str(text)


def make_bar_chart(df, iv, dv):
    """Grouped bar chart styled to match native SPSS output: light-gray plot
    area, black axis border, SPSS categorical palette, boxed % data labels,
    and the DV question wrapped as the legend title."""
    sub = df[[iv, dv]].dropna()
    if sub.empty or sub[iv].nunique() < 2 or sub[dv].nunique() < 2:
        return None, None
    # Hard safety ceiling regardless of how the column was classified --
    # a chart with dozens of x-axis groups or legend entries isn't
    # meaningful anyway, and rendering one can spike memory/CPU sharply.
    if sub[iv].nunique() > 15 or sub[dv].nunique() > 10:
        return None, None

    iv_order = _ordered_categories(sub[iv], prefer_likert=False)
    dv_order = _ordered_categories(sub[dv], prefer_likert=True)

    sub = sub.copy()
    sub[iv] = pd.Categorical(sub[iv], categories=iv_order, ordered=True)
    sub[dv] = pd.Categorical(sub[dv], categories=dv_order, ordered=True)

    ct = pd.crosstab(sub[iv], sub[dv], normalize="index") * 100
    ct = ct.reindex(columns=dv_order)
    ct = ct.round(2)

    n_cats = len(ct.columns)
    fig, ax = plt.subplots(figsize=(6.2, 4.6), dpi=80)
    try:
        fig.patch.set_facecolor("white")
        ax.set_facecolor(SPSS_BG)

        x = np.arange(len(ct.index))
        n_groups = len(ct.index)
        width = 0.8 / max(n_cats, 1)

        for i, cat in enumerate(ct.columns):
            vals = ct[cat].values
            offset = (i - (n_cats - 1) / 2) * width
            bars = ax.bar(
                x + offset, vals, width=width,
                color=SPSS_PALETTE[i % len(SPSS_PALETTE)],
                edgecolor="black", linewidth=0.6, label=str(cat),
            )
            for rect, v in zip(bars, vals):
                if v <= 0:
                    continue
                ax.annotate(
                    f"{v:.2f}%",
                    xy=(rect.get_x() + rect.get_width() / 2, v),
                    xytext=(0, 3), textcoords="offset points",
                    ha="center", va="bottom", fontsize=7,
                    bbox=dict(boxstyle="square,pad=0.25", fc="white", ec="black", lw=0.6),
                )

        ax.set_xticks(x)
        ax.set_xticklabels([str(c) for c in ct.index], fontsize=9, fontweight="bold")
        ax.set_ylabel("Percent", fontsize=10, fontweight="bold")
        ax.set_xlabel(str(iv), fontsize=10, fontweight="bold")
        ymax = max(ct.values.max() * 1.25, 10) if ct.values.size else 10
        ax.set_ylim(0, ymax)

        for spine in ax.spines.values():
            spine.set_visible(True)
            spine.set_color("black")
            spine.set_linewidth(0.8)

        legend = ax.legend(
            title=_wrap_label(dv, width=16),
            fontsize=8, title_fontsize=8,
            bbox_to_anchor=(1.02, 1), loc="upper left",
            frameon=False, borderaxespad=0,
        )
        legend.get_title().set_ha("left")

        plt.tight_layout()

        buf = io.BytesIO()
        fig.savefig(buf, format="png", bbox_inches="tight", facecolor="white")
        buf.seek(0)
        return buf.getvalue(), ct
    finally:
        # Always release the figure, even if plotting raised partway through --
        # otherwise a single bad column can leak a full Figure per request.
        plt.close(fig)


def auto_legend_text(iv, dv, ct):
    """Rule-based legend/result sentence — the offline fallback, always available."""
    try:
        top_group = ct.max(axis=1).idxmax()
        top_cat = ct.loc[top_group].idxmax()
        top_pct = ct.loc[top_group, top_cat]
        return (
            f"The given figure represents the {iv}-wise distribution of respondents and their "
            f"views on \u201c{dv}\u201d. The {top_group} group shows the strongest single pattern: "
            f"{top_pct:.0f}% selected \u201c{top_cat}\u201d, higher than any other group/category "
            f"combination in this comparison."
        )
    except Exception:
        return f"The chart compares {dv} across {iv} groups."


def auto_result_sentence(iv, dv, ct):
    """Rule-based, offline 'Results'-style sentence for one chart (objective, factual)."""
    try:
        top_group = ct.max(axis=1).idxmax()
        top_cat = ct.loc[top_group].idxmax()
        top_pct = ct.loc[top_group, top_cat]
        return (
            f"According to the chart, {top_pct:.0f}% of {top_group} respondents selected "
            f"\u201c{top_cat}\u201d for \u201c{dv}\u201d, the strongest single pattern across {iv} groups."
        )
    except Exception:
        return f"The chart reports responses to \u201c{dv}\u201d across {iv} groups."


def auto_discussion_sentence(iv, dv, ct):
    """Rule-based, offline 'Discussion'-style sentence for one chart (interpretive)."""
    try:
        top_group = ct.max(axis=1).idxmax()
        top_cat = ct.loc[top_group].idxmax()
        top_pct = ct.loc[top_group, top_cat]
        return (
            f"This suggests {top_group} respondents feel most strongly about \u201c{dv}\u201d, "
            f"with {top_pct:.0f}% converging on \u201c{top_cat}\u201d -- a pattern worth weighing "
            f"against the other {iv} groups shown."
        )
    except Exception:
        return f"This points to a notable pattern in how {iv} groups responded to \u201c{dv}\u201d."


def chart_fact(iv, dv, ct):
    """Compact numeric summary of one chart, used as input to the AI writer."""
    try:
        top_group = ct.max(axis=1).idxmax()
        top_cat = ct.loc[top_group].idxmax()
        top_pct = float(ct.loc[top_group, top_cat])
        return {"iv": iv, "dv": dv, "top_group": str(top_group), "top_category": str(top_cat), "top_pct": round(top_pct, 1)}
    except Exception:
        return {"iv": iv, "dv": dv, "top_group": None, "top_category": None, "top_pct": None}


def ai_generate_narrative(facts, anova, composite_label):
    """
    One batched call to Groq's free-tier API that writes every chart legend
    plus one Results-style and one Discussion-style sentence per chart (in
    the same "ADDITIONAL ANALYSIS" report style used at this org: a running
    narrative that cites each chart as it goes, e.g. "...(fig: 3)"). Keeps
    this well within free rate limits even for 50+ charts since it's one
    batched request. Returns None on any failure so the caller falls back
    to the rule-based text -- the (fig: N) citations themselves are always
    appended in code afterward, not trusted to the model, so numbering
    never drifts even if the model's prose does.
    """
    if _groq_client is None:
        return None

    anova_summary = None
    if anova is not None:
        anova_summary = {
            "grouping_variable": anova["iv"],
            "outcome": composite_label,
            "F": round(anova["F"], 3),
            "p": round(anova["p"], 4),
            "significant": anova["significant"],
            "df_between": anova["df_between"],
            "df_within": anova["df_within"],
        }

    prompt = (
        "You are writing an academic-style SPSS survey report in the following house style: "
        "a numbered FIGURE/LEGEND block per chart, followed by a RESULT section and a DISCUSSION "
        "section that are each ONE continuous narrative walking through every chart in order "
        "(each chart gets 1-3 sentences citing concrete percentages/groups, then the narrative "
        "moves to the next chart) -- NOT a short high-level summary paragraph. "
        "Given the JSON facts below, respond with ONLY valid JSON (no markdown, no code fences) "
        "matching this exact shape:\n"
        '{"legends": ["...", ...], "results_sentences": ["...", ...], "discussion_sentences": ["...", ...]}\n\n'
        "Rules:\n"
        "- \"legends\" must have exactly one string per item in chart_facts, in the same order, "
        "each 1-2 sentences describing that specific chart's pattern (mention the top group/category/percent).\n"
        "- \"results_sentences\" must ALSO have exactly one string per item in chart_facts, in the same order: "
        "an objective 1-3 sentence description of that chart's data (concrete percentages/groups). "
        "Do not add a figure citation yourself -- it is appended automatically afterward.\n"
        "- \"discussion_sentences\" must ALSO have exactly one string per item in chart_facts, in the same order: "
        "a more interpretive 1-3 sentence take on what that chart's pattern suggests. Do not add a citation.\n"
        "- If anova_summary is present, make the very last discussion_sentences item also mention the ANOVA "
        "F, df, and p value in APA style and what it means for significance.\n"
        "- Do not invent numbers not present in the facts.\n\n"
        f"chart_facts = {json.dumps(facts)}\n"
        f"anova_summary = {json.dumps(anova_summary)}\n"
    )

    try:
        resp = _groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": "You write precise, concise, factual survey-analysis reports. Output strict JSON only."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.4,
            max_tokens=4000,
            response_format={"type": "json_object"},
            timeout=20,
        )
        text = resp.choices[0].message.content
        data = json.loads(text)
        if (
            isinstance(data.get("legends"), list) and len(data["legends"]) == len(facts)
            and isinstance(data.get("results_sentences"), list) and len(data["results_sentences"]) == len(facts)
            and isinstance(data.get("discussion_sentences"), list) and len(data["discussion_sentences"]) == len(facts)
        ):
            return data
        return None
    except Exception:
        traceback.print_exc()
        return None


# ---------------------------------------------------------------------------
# ANOVA (one-way, SPSS-style table)
# ---------------------------------------------------------------------------

def run_anova(df, iv, dv_numeric_col_name, dv_values):
    work = pd.DataFrame({iv: df[iv], "dv": dv_values}).dropna()
    groups = [g["dv"].values for _, g in work.groupby(iv) if len(g) > 1]
    labels = [name for name, g in work.groupby(iv) if len(g) > 1]
    if len(groups) < 2:
        return None

    grand_mean = work["dv"].mean()
    n_total = len(work)
    k = len(groups)

    ss_between = sum(len(g) * (g.mean() - grand_mean) ** 2 for g in groups)
    ss_within = sum(((g - g.mean()) ** 2).sum() for g in groups)
    ss_total = ss_between + ss_within

    df_between = k - 1
    df_within = n_total - k

    ms_between = ss_between / df_between if df_between else float("nan")
    ms_within = ss_within / df_within if df_within else float("nan")
    F = ms_between / ms_within if ms_within else float("nan")
    p = stats.f.sf(F, df_between, df_within) if df_within > 0 else float("nan")

    group_stats = work.groupby(iv)["dv"].agg(["count", "mean", "std"]).round(2)

    return {
        "iv": iv,
        "dv_label": dv_numeric_col_name,
        "ss_between": ss_between, "ss_within": ss_within, "ss_total": ss_total,
        "df_between": df_between, "df_within": df_within,
        "ms_between": ms_between, "ms_within": ms_within,
        "F": F, "p": p,
        "group_stats": group_stats,
        "significant": bool(p < 0.05) if p == p else False,
    }


# ---------------------------------------------------------------------------
# SPSS-style statistical output tables, rendered as images (matplotlib table)
# so they read as authentic SPSS "export table as image" output rather than
# a native Word table.
# ---------------------------------------------------------------------------

def _spss_table_image(caption, col_labels, rows, footnote=None, col_widths=None):
    """Render a classic SPSS output table (bold caption, thin horizontal
    rules, no vertical grid lines, small italic footnote) to a PNG."""
    n_rows = len(rows) + 1  # + header
    n_cols = len(col_labels)
    fig_w = max(4.2, 1.15 * n_cols)
    fig_h = 0.42 * n_rows + 0.7
    fig, ax = plt.subplots(figsize=(fig_w, fig_h), dpi=110)
    ax.axis("off")

    table = ax.table(
        cellText=rows,
        colLabels=col_labels,
        cellLoc="center",
        colLoc="center",
        loc="center",
        colWidths=col_widths,
    )
    table.auto_set_font_size(False)
    table.set_fontsize(8.5)
    table.scale(1, 1.6)

    for (r, c), cell in table.get_celld().items():
        cell.set_edgecolor("white")
        cell.set_linewidth(0)
        cell.set_text_props(color="black")
        if r == 0:
            cell.set_text_props(fontweight="bold")
            cell.visible_edges = "TB"
        elif r == n_rows - 1:
            cell.visible_edges = "B"
        else:
            cell.visible_edges = ""
        cell.set_edgecolor("black")

    ax.set_title(caption, fontsize=10, fontweight="bold", pad=14)

    if footnote:
        fig.text(0.02, 0.02, footnote, fontsize=6.5, style="italic", ha="left", va="bottom")

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", facecolor="white", dpi=110)
    buf.seek(0)
    data = buf.getvalue()
    plt.close(fig)
    return data


def run_chi_square(df, iv, dv):
    """Pearson chi-square test of independence on the iv x dv crosstab,
    SPSS 'Chi-Square Tests' style (Pearson, Likelihood Ratio, Linear-by-Linear)."""
    sub = df[[iv, dv]].dropna()
    if sub.empty or sub[iv].nunique() < 2 or sub[dv].nunique() < 2:
        return None
    ct = pd.crosstab(sub[iv], sub[dv])
    if ct.shape[0] < 2 or ct.shape[1] < 2:
        return None
    try:
        chi2, p, dof, expected = stats.chi2_contingency(ct)
    except Exception:
        return None

    # Likelihood ratio G-test (same df as Pearson chi-square)
    try:
        g_stat, g_p, _, _ = stats.chi2_contingency(ct, lambda_="log-likelihood")
    except Exception:
        g_stat, g_p = float("nan"), float("nan")

    n_expected_low = int((expected < 5).sum())
    pct_low = 100 * n_expected_low / expected.size if expected.size else 0
    min_expected = float(expected.min()) if expected.size else float("nan")

    return {
        "iv": iv, "dv": dv,
        "chi2": chi2, "p": p, "dof": dof,
        "g_stat": g_stat, "g_p": g_p,
        "n_valid": int(ct.values.sum()),
        "n_expected_low": n_expected_low, "pct_expected_low": pct_low,
        "min_expected": min_expected,
    }


def chi_square_table_image(cs):
    def fmt_p(p):
        return "<.001" if p < 0.001 else f"{p:.3f}"
    rows = [
        ["Pearson Chi-Square", f"{cs['chi2']:.3f}", str(cs["dof"]), fmt_p(cs["p"])],
        ["Likelihood Ratio", f"{cs['g_stat']:.3f}", str(cs["dof"]), fmt_p(cs["g_p"])],
        ["N of Valid Cases", "", "", str(cs["n_valid"])],
    ]
    footnote = (
        f"{cs['n_expected_low']} cells ({cs['pct_expected_low']:.1f}%) have expected count "
        f"less than 5. The minimum expected count is {cs['min_expected']:.2f}."
    )
    return _spss_table_image(
        "Chi-Square Tests",
        ["", "Value", "df", "Asymptotic Significance\n(2-sided)"],
        rows, footnote=footnote, col_widths=[0.34, 0.18, 0.14, 0.34],
    )


def anova_table_image(anova):
    def fmt_p(p):
        return "<.001" if p < 0.001 else f"{p:.3f}"
    rows = [
        ["Between Groups", f"{anova['ss_between']:.3f}", str(anova["df_between"]),
         f"{anova['ms_between']:.3f}", f"{anova['F']:.3f}", fmt_p(anova["p"])],
        ["Within Groups", f"{anova['ss_within']:.3f}", str(anova["df_within"]),
         f"{anova['ms_within']:.3f}", "", ""],
        ["Total", f"{anova['ss_total']:.3f}", str(anova["df_between"] + anova["df_within"]), "", "", ""],
    ]
    return _spss_table_image(
        "ANOVA",
        ["", "Sum of Squares", "df", "Mean Square", "F", "Sig."],
        rows, col_widths=[0.22, 0.18, 0.1, 0.18, 0.14, 0.14],
    )


# ---------------------------------------------------------------------------
# Word report builder
# ---------------------------------------------------------------------------

def _set_base_font(doc):
    """Match the reference report's font: Times New Roman 12pt body text."""
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)


def _labeled_paragraph(doc, label, text):
    """A paragraph like '**LEGEND:** some text', matching the reference style."""
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(text)
    return p


def build_docx(title, iv_cols, dv_cols, chart_records, anova, composite_label,
                ai_results_sentences=None, ai_discussion_sentences=None, ai_used=False,
                chi_square=None, charts_capped=False, total_possible=None):
    doc = Document()
    _set_base_font(doc)

    # Title -- bold plain paragraph, not a big Word Heading style.
    title_p = doc.add_paragraph()
    title_p.add_run("ADDITIONAL ANALYSIS").bold = True
    doc.add_paragraph()

    if charts_capped:
        note_p = doc.add_paragraph()
        note_run = note_p.add_run(
            f"Note: This spreadsheet had {total_possible} possible variable-pair charts. "
            f"Only the first {len(chart_records)} are included below to keep the report "
            f"within server memory limits. Set the MAX_CHARTS environment variable higher "
            f"if you're self-hosting and need every combination."
        )
        note_run.italic = True
        doc.add_paragraph()

    fig_num = 0

    # One FIGURE / image / LEGEND block per chart, in order.
    for rec in chart_records:
        fig_num += 1
        fig_p = doc.add_paragraph()
        fig_p.add_run(f"FIGURE {fig_num}:").bold = True

        img_stream = io.BytesIO(rec["png"])
        doc.add_picture(img_stream, width=Inches(3.15))

        _labeled_paragraph(doc, "LEGEND", rec["legend"])
        doc.add_paragraph()

    # Chi-Square Tests, presented as its own numbered figure -- an authentic
    # SPSS "Chi-Square Tests" output table (rendered as an image, not a
    # native Word table) for the primary IV x DV crosstab.
    if chi_square is not None:
        fig_num += 1
        fig_p = doc.add_paragraph()
        fig_p.add_run(f"FIGURE {fig_num}:").bold = True
        doc.add_picture(io.BytesIO(chi_square_table_image(chi_square)), width=Inches(3.6))

        cs_sig = chi_square["p"] < 0.05
        cs_p_txt = "<.001" if chi_square["p"] < 0.001 else f"{chi_square['p']:.3f}"
        _labeled_paragraph(
            doc, "LEGEND",
            f"This set of statistical outputs of Chi-Square presents a breakdown of respondent "
            f"engagement and opinion across {chi_square['iv']} groups regarding {chi_square['dv']}."
        )
        _labeled_paragraph(
            doc, "INTERPRETATION",
            f"The Pearson Chi-Square value is {chi_square['chi2']:.3f} with {chi_square['dof']} "
            f"degree(s) of freedom and a significance level of {cs_p_txt}, which is "
            f"{'below' if cs_sig else 'above'} the conventional 0.05 threshold. This indicates "
            f"that {chi_square['iv']} and {chi_square['dv']} "
            f"{'are' if cs_sig else 'are not'} significantly associated, based on "
            f"{chi_square['n_valid']} valid cases."
        )
        doc.add_paragraph()

    # ANOVA presented as its own numbered figure -- an authentic SPSS "ANOVA"
    # output table (image), not a native Word table.
    if anova is not None:
        fig_num += 1
        fig_p = doc.add_paragraph()
        fig_p.add_run(f"FIGURE {fig_num}:").bold = True
        doc.add_picture(io.BytesIO(anova_table_image(anova)), width=Inches(4.0))
        doc.add_paragraph()

        sig_txt = (
            f"The model shows F({anova['df_between']}, {anova['df_within']}) = {anova['F']:.3f}, "
            f"p = {anova['p']:.3f}, which is {'below' if anova['significant'] else 'above'} the "
            f"conventional 0.05 threshold. This indicates that {anova['iv']} "
            f"{'does' if anova['significant'] else 'does not'} have a statistically significant "
            f"effect on {composite_label}."
        )
        _labeled_paragraph(
            doc, "LEGEND",
            f"This ANOVA table evaluates whether {anova['iv']} significantly predicts {composite_label}."
        )
        _labeled_paragraph(doc, "INTERPRETATION", sig_txt)
        doc.add_paragraph()

    # RESULT: one continuous narrative, one entry per chart, each citing (fig: N).
    result_label_p = doc.add_paragraph()
    result_label_p.add_run("RESULT:").bold = True
    result_p = doc.add_paragraph()
    for idx, rec in enumerate(chart_records, start=1):
        sentence = (
            ai_results_sentences[idx - 1] if ai_results_sentences and idx - 1 < len(ai_results_sentences)
            else auto_result_sentence(rec["iv"], rec["dv"], rec["ct"])
        )
        result_p.add_run(sentence.strip() + " ")
        result_p.add_run(f"(fig: {idx}) ").bold = True
    doc.add_paragraph()

    # DISCUSSION: same shape, more interpretive tone, plus an ANOVA callout at the end.
    disc_label_p = doc.add_paragraph()
    disc_label_p.add_run("DISCUSSION").bold = True
    disc_p = doc.add_paragraph()
    for idx, rec in enumerate(chart_records, start=1):
        sentence = (
            ai_discussion_sentences[idx - 1] if ai_discussion_sentences and idx - 1 < len(ai_discussion_sentences)
            else auto_discussion_sentence(rec["iv"], rec["dv"], rec["ct"])
        )
        disc_p.add_run(sentence.strip() + " ")
        disc_p.add_run(f"(fig: {idx}) ").bold = True
    if anova is not None:
        closing = (
            f"Taken together with the ANOVA above (F = {anova['F']:.3f}, p = {anova['p']:.3f}), "
            f"{anova['iv']} {'does' if anova['significant'] else 'does not'} appear to have a "
            f"statistically significant relationship with {composite_label}, so any differences "
            f"seen across the charts should be weighed against that result and the usual caveats "
            f"of survey data (sample size, response bias, and measurement limitations)."
        )
        disc_p.add_run(closing)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/analyze", methods=["POST"])
def analyze():
    try:
        f = request.files.get("file")
        if not f:
            return jsonify({"error": "No file uploaded."}), 400

        # Cap extremely large spreadsheets *before* fully loading them --
        # capping rows only after pd.read_csv/read_excel has already parsed
        # the entire file into memory doesn't help; a large/wide file can
        # OOM the 512MB Render free instance during the read itself,
        # especially .xlsx (openpyxl commonly uses 10-20x the file's raw
        # size in RAM). Reading with nrows keeps peak memory bounded by the
        # cap, not by the uploaded file's size.
        MAX_ROWS = int(os.environ.get("MAX_ROWS", 1500))

        filename = f.filename.lower()
        if filename.endswith(".csv"):
            df = pd.read_csv(f, nrows=MAX_ROWS)
        else:
            df = pd.read_excel(f, nrows=MAX_ROWS)

        df = df.dropna(axis=1, how="all")
        if df.empty:
            return jsonify({"error": "The file has no usable data."}), 400

        # Clean up messy header text (stray newlines/whitespace from form exports)
        df.columns = [" ".join(str(c).split()) for c in df.columns]

        iv_cols, dv_cols = classify_columns(df)
        if not iv_cols or not dv_cols:
            return jsonify({"error": "Could not detect independent/dependent variables automatically."}), 400

        # Generate every IV x DV combination by default (matches real SPSS
        # exports, which don't sample a subset of charts) -- but cap the
        # total so a wide/long spreadsheet can't OOM-kill a free-tier Render
        # worker (512MB RAM). Unlike the old silent 6-chart cap, we tell the
        # caller (and the user, via the UI) whenever the cap actually bites.
        MAX_CHARTS = int(os.environ.get("MAX_CHARTS", 60))
        chart_records = []
        preview_records = []
        facts = []
        count = 0
        total_possible = 0
        for dv in dv_cols:
            for iv in iv_cols:
                total_possible += 1
                if count >= MAX_CHARTS:
                    continue
                png, ct = make_bar_chart(df, iv, dv)
                if png is None:
                    continue
                fallback_legend = auto_legend_text(iv, dv, ct)
                chart_records.append({"iv": iv, "dv": dv, "png": png, "legend": fallback_legend, "ct": ct})
                facts.append(chart_fact(iv, dv, ct))
                count += 1
                if count % 20 == 0:
                    # Release matplotlib/pandas intermediates periodically so
                    # memory doesn't climb monotonically across a long loop.
                    gc.collect()
        charts_capped = total_possible > count

        # Build a composite DV score (mean of all Likert-encoded DVs) for the ANOVA
        numeric_cols = []
        for dv in dv_cols:
            numeric_series, is_scale = to_numeric_dv(df, dv)
            if is_scale:
                numeric_cols.append(numeric_series.rename(dv))
        anova = None
        composite_label = "overall attitude score"
        primary_iv = iv_cols[0]
        if numeric_cols:
            composite = pd.concat(numeric_cols, axis=1).mean(axis=1)
            anova = run_anova(df, primary_iv, composite_label, composite)

        # Chi-Square Tests figure: primary IV x primary (first) DV crosstab,
        # matching the "Chi-Square Tests" output block seen in real SPSS reports.
        chi_square = None
        if dv_cols:
            chi_square = run_chi_square(df, primary_iv, dv_cols[0])

        # Try one batched Groq call to write every legend + per-chart Results/Discussion
        # sentences. Falls back silently (ai_used=False) if no key is set or the call fails.
        ai_used = False
        ai_results_sentences = None
        ai_discussion_sentences = None
        ai_data = ai_generate_narrative(facts, anova, composite_label)
        if ai_data:
            ai_used = True
            for rec, legend in zip(chart_records, ai_data["legends"]):
                rec["legend"] = legend
            ai_results_sentences = ai_data["results_sentences"]
            ai_discussion_sentences = ai_data["discussion_sentences"]

        # Build the DOCX once during analysis. Rebuilding all charts inside
        # /download could exceed Render's request timeout and return an HTML 500 page.
        title = "GraphGen Pro - Automated Analysis Report"
        report_docx = build_docx(
            title, iv_cols, dv_cols, chart_records, anova, composite_label,
            ai_results_sentences=ai_results_sentences,
            ai_discussion_sentences=ai_discussion_sentences,
            ai_used=ai_used,
            chi_square=chi_square,
            charts_capped=charts_capped,
            total_possible=total_possible,
        )
        report_bytes = report_docx.getvalue()
        del report_docx  # the in-memory Document + all its embedded images

        report_id = str(uuid.uuid4())
        REPORT_CACHE[report_id] = {
            "docx_bytes": report_bytes,
        }
        if len(REPORT_CACHE) > MAX_CACHE:
            REPORT_CACHE.pop(next(iter(REPORT_CACHE)))

        # Drop the (potentially large) source dataframe now that everything
        # needed from it has been extracted, and force a GC pass so this
        # worker's memory footprint doesn't keep climbing across requests.
        n_rows = len(df)
        del df
        gc.collect()

        # Only send a small preview to the browser; full images stay server-side.
        for rec in chart_records[:6]:
            preview_records.append({
                "iv": rec["iv"], "dv": rec["dv"], "legend": rec["legend"],
                "img": "data:image/png;base64," + base64.b64encode(rec["png"]).decode(),
            })

        anova_summary = None
        if anova is not None:
            anova_summary = {
                "iv": anova["iv"], "F": round(anova["F"], 3), "p": round(anova["p"], 3),
                "significant": anova["significant"],
            }

        return jsonify({
            "report_id": report_id,
            "n_rows": int(n_rows),
            "iv_cols": iv_cols, "dv_cols": dv_cols,
            "total_charts": len(chart_records),
            "total_possible_charts": total_possible,
            "charts_capped": charts_capped,
            "preview_charts": preview_records,
            "anova": anova_summary,
            "ai_used": ai_used,
        })

    except MemoryError:
        traceback.print_exc()
        gc.collect()
        return jsonify({
            "error": "The file is too large/complex to process on this server's memory limit. "
                     "Try a smaller file, fewer columns, or lower MAX_ROWS/MAX_CHARTS."
        }), 500
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"Something went wrong: {e}"}), 500


@app.route("/download/<report_id>")
def download(report_id):
    data = REPORT_CACHE.get(report_id)
    if not data:
        return jsonify({
            "error": "Report expired or not found. Please re-analyze your file."
        }), 404

    # Lightweight endpoint: the DOCX was already generated during /analyze.
    docx_bytes = data.get("docx_bytes")
    if not docx_bytes:
        return jsonify({
            "error": "The report file is unavailable. Please re-analyze your file."
        }), 500

    return send_file(
        io.BytesIO(docx_bytes),
        as_attachment=True,
        download_name="GraphGen_Pro_Report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        max_age=0,
    )


@app.errorhandler(413)
def request_too_large(error):
    return jsonify({
        "error": "File is too large. Please upload an XLSX/CSV file smaller than 8 MB."
    }), 413


@app.errorhandler(Exception)
def handle_unexpected_error(error):
    # API routes return JSON instead of Flask/Render's HTML error page.
    traceback.print_exc()
    if request.path.startswith("/analyze") or request.path.startswith("/download/"):
        return jsonify({"error": "Internal server error. Check the server logs for details."}), 500
    return error


@app.route("/health")
def health():
    return "ok"


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
